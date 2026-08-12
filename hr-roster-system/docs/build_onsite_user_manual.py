from __future__ import annotations

from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/zhuhao/Documents/moluo/hr-roster-system")
OUT_DIR = ROOT / "docs" / "驻厂专员使用手册"
IMG_DIR = OUT_DIR / "配图"
DOCX_PATH = OUT_DIR / "优益数字化管理系统_驻厂专员使用手册_V1.0.docx"

FONT_PATH = "/System/Library/Fonts/Hiragino Sans GB.ttc"
FONT_FALLBACK = "/System/Library/Fonts/STHeiti Medium.ttc"

NAVY = "071B33"
NAVY_2 = "0B2948"
TEAL = "0EA5A8"
CYAN = "35C7D0"
BLUE = "2563EB"
LIGHT_BLUE = "EAF7FA"
PALE = "F4F8FB"
GREEN = "16A36A"
AMBER = "F59E0B"
RED = "DC4C4C"
GRAY = "617184"
LIGHT_GRAY = "DDE6EE"
TEXT = "182739"
WHITE = "FFFFFF"


def font(size: int, bold: bool = False):
    path = FONT_PATH if Path(FONT_PATH).exists() else FONT_FALLBACK
    index = 1 if bold else 0
    try:
        return ImageFont.truetype(path, size=size, index=index)
    except Exception:
        return ImageFont.truetype(path, size=size)


def rounded(draw, box, radius=24, fill="#FFFFFF", outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def txt(draw, xy, content, size=28, fill="#182739", bold=False, anchor=None):
    draw.text(xy, content, font=font(size, bold), fill=fill, anchor=anchor)


def fit_text(draw, content: str, max_width: int, size=28, bold=False):
    current = size
    while current > 16 and draw.textbbox((0, 0), content, font=font(current, bold))[2] > max_width:
        current -= 1
    return font(current, bold)


def arrow(draw, start, end, color="#0EA5A8", width=8):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - sign * 22, y2 - 15), (x2 - sign * 22, y2 + 15)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 15, y2 - sign * 22), (x2 + 15, y2 - sign * 22)]
    draw.polygon(points, fill=color)


def canvas(title: str, subtitle: str = ""):
    image = Image.new("RGB", (1600, 900), "#EDF4F8")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 128), fill="#071B33")
    draw.rectangle((0, 124, 1600, 128), fill="#20C5C8")
    txt(draw, (72, 42), title, 44, "#FFFFFF", True)
    if subtitle:
        txt(draw, (72, 94), subtitle, 21, "#9FDDE1")
    return image, draw


def save(image: Image.Image, name: str):
    path = IMG_DIR / name
    image.save(path, quality=95)
    return path


def make_cover():
    image = Image.new("RGB", (1600, 1050), "#071B33")
    draw = ImageDraw.Draw(image)
    for i in range(14):
        x = 860 + i * 72
        draw.line((x, 0, x - 530, 1050), fill="#0B2948", width=2)
    for i in range(12):
        y = 100 + i * 76
        draw.line((760, y, 1600, y), fill="#0D3856", width=2)
    draw.ellipse((1080, 105, 1510, 535), fill="#0B3850", outline="#23C4C8", width=4)
    draw.ellipse((1180, 205, 1410, 435), fill="#0EA5A8")
    txt(draw, (1295, 320), "优", 110, "#FFFFFF", True, "mm")

    txt(draw, (96, 118), "YOUYI DIGITAL HR", 25, "#58D5D8", True)
    txt(draw, (96, 206), "优益数字化管理系统", 58, "#FFFFFF", True)
    txt(draw, (96, 294), "驻厂专员使用手册", 76, "#FFFFFF", True)
    txt(draw, (100, 400), "员工入职 · 在职维护 · 合同雇主险 · 预支 · 离职", 31, "#B8D4E4")

    rounded(draw, (96, 500, 700, 785), 30, "#0B2948", "#1F6078", 2)
    items = [
        ("01", "按客户/项目管理员工"),
        ("02", "现场操作全流程留痕"),
        ("03", "只查看本人授权范围"),
    ]
    for idx, (num, label) in enumerate(items):
        y = 555 + idx * 76
        rounded(draw, (130, y - 15, 190, y + 45), 16, "#0EA5A8")
        txt(draw, (160, y + 14), num, 22, "#FFFFFF", True, "mm")
        txt(draw, (220, y + 15), label, 29, "#FFFFFF", idx == 0, "lm")

    # 手机界面示意
    rounded(draw, (850, 430, 1260, 960), 48, "#FFFFFF", "#2CC7CA", 5)
    rounded(draw, (875, 465, 1235, 570), 24, "#071B33")
    txt(draw, (900, 492), "FIELD OPERATIONS", 18, "#70DADD", True)
    txt(draw, (900, 532), "驻厂工作台", 33, "#FFFFFF", True)
    for i, (label, value, color) in enumerate([
        ("在职", "86", "#0EA5A8"), ("待入职", "12", "#2563EB"), ("待处理", "5", "#F59E0B")
    ]):
        x = 895 + i * 108
        rounded(draw, (x, 600, x + 92, 685), 18, "#EDF6F8")
        txt(draw, (x + 46, 625), label, 16, "#617184", False, "mm")
        txt(draw, (x + 46, 660), value, 28, color, True, "mm")
    for i, label in enumerate(["新增员工", "我的员工", "雇主险", "驻厂预支"]):
        x = 895 + (i % 2) * 170
        y = 720 + (i // 2) * 100
        rounded(draw, (x, y, x + 150, y + 78), 18, "#F4F8FB", "#D5E6ED")
        txt(draw, (x + 75, y + 40), label, 21, "#18334D", True, "mm")
    txt(draw, (98, 970), "版本 V1.0  |  适用：企业管理员授权的驻厂专员账号", 23, "#8EB2C7")
    return save(image, "01_封面.png")


def make_permission():
    image, draw = canvas("驻厂专员的数据权限", "企业管理员先派遣项目，驻厂账号才能看到对应客户、项目和员工")
    cards = [
        (80, 250, 390, 500, "企业管理员", "建立客户与项目\n选择“派遣驻厂”\n指定驻厂专员", "#2563EB"),
        (645, 250, 955, 500, "项目授权", "系统建立账号与\n客户/项目的授权关系", "#0EA5A8"),
        (1210, 250, 1520, 500, "驻厂专员", "只查看授权项目\n维护本项目员工\n处理现场事项", "#16A36A"),
    ]
    for x1, y1, x2, y2, title, body, color in cards:
        rounded(draw, (x1, y1, x2, y2), 26, "#FFFFFF", color, 3)
        draw.ellipse((x1 + 110, y1 + 35, x1 + 200, y1 + 125), fill=color)
        txt(draw, ((x1 + x2) // 2, y1 + 170), title, 31, "#182739", True, "mm")
        for n, line in enumerate(body.split("\n")):
            txt(draw, ((x1 + x2) // 2, y1 + 215 + n * 38), line, 23, "#617184", False, "mm")
    arrow(draw, (410, 375), (620, 375))
    arrow(draw, (975, 375), (1185, 375))
    rounded(draw, (250, 650, 1350, 780), 26, "#FFF7E7", "#F59E0B", 2)
    txt(draw, (310, 695), "权限边界", 27, "#B56A00", True)
    txt(draw, (310, 742), "看不到其他项目属于正常的数据隔离；需要新增范围时，由企业管理员重新派遣。", 27, "#694C1F")
    return save(image, "02_驻厂权限关系图.png")


def make_flow():
    image, draw = canvas("驻厂员工全生命周期", "从招聘录入到离职回流人才库，关键操作均保留记录")
    nodes = [
        ("1", "录入员工", "新增/批量/OCR", "#2563EB"),
        ("2", "待入职", "核验客户岗位", "#0EA5A8"),
        ("3", "确认入职", "合同+雇主险", "#16A36A"),
        ("4", "在职维护", "编辑/调岗/预支", "#6D5BD0"),
        ("5", "离职交接", "减保+工资结算", "#F59E0B"),
        ("6", "人才库", "离职/未入职回流", "#DC4C4C"),
    ]
    for i, (num, title, sub, color) in enumerate(nodes):
        x = 70 + i * 255
        rounded(draw, (x, 300, x + 205, 520), 28, "#FFFFFF", color, 3)
        draw.ellipse((x + 65, 330, x + 140, 405), fill=color)
        txt(draw, (x + 102, 368), num, 29, "#FFFFFF", True, "mm")
        txt(draw, (x + 102, 447), title, 28, "#182739", True, "mm")
        txt(draw, (x + 102, 487), sub, 19, "#617184", False, "mm")
        if i < len(nodes) - 1:
            arrow(draw, (x + 210, 410), (x + 245, 410), color="#3AAFB5", width=6)
    rounded(draw, (180, 650, 1420, 786), 28, "#EAF7FA", "#0EA5A8", 2)
    txt(draw, (230, 690), "入职风险只检查两项", 27, "#087D80", True)
    txt(draw, (230, 740), "① 劳动合同是否签订　　② 雇主险是否增保", 32, "#182739", True)
    return save(image, "03_员工生命周期流程图.png")


def mobile_frame(draw, title, subtitle="", x=460, y=160, w=680, h=690):
    rounded(draw, (x, y, x + w, y + h), 42, "#FFFFFF", "#B6D2DE", 3)
    rounded(draw, (x + 20, y + 20, x + w - 20, y + 125), 24, "#071B33")
    txt(draw, (x + 52, y + 52), title, 34, "#FFFFFF", True)
    if subtitle:
        txt(draw, (x + 52, y + 94), subtitle, 19, "#8DDADC")
    return x + 35, y + 150, w - 70, h - 180


def make_workbench():
    image, draw = canvas("小程序：驻厂工作台", "首页聚合人员、项目、预支、合规待办和现场快捷入口")
    x, y, w, h = mobile_frame(draw, "驻厂工作台", "现场运营在线")
    metrics = [("在职人员", "86"), ("在营项目", "3"), ("未结预支", "¥2,300")]
    for i, (label, value) in enumerate(metrics):
        bx = x + i * 198
        rounded(draw, (bx, y, bx + 178, y + 92), 18, "#EEF7F9")
        txt(draw, (bx + 89, y + 29), label, 18, "#617184", False, "mm")
        txt(draw, (bx + 89, y + 66), value, 27, "#0A8F93", True, "mm")
    txt(draw, (x, y + 132), "驻厂人员管理", 25, "#182739", True)
    actions = [
        ("01", "录入新员工", "先录档案，再确认到岗"),
        ("02", "待入职人员", "核验身份、客户和岗位"),
        ("03", "在职员工", "资料、调岗与雇主险"),
        ("04", "离职交接", "减保、交接和工资结算"),
    ]
    for i, (num, title, sub) in enumerate(actions):
        bx = x + (i % 2) * 300
        by = y + 170 + (i // 2) * 128
        rounded(draw, (bx, by, bx + 280, by + 105), 18, "#FFFFFF", "#D9E8EE")
        rounded(draw, (bx + 18, by + 18, bx + 68, by + 68), 13, "#0EA5A8")
        txt(draw, (bx + 43, by + 43), num, 17, "#FFFFFF", True, "mm")
        txt(draw, (bx + 82, by + 30), title, 21, "#182739", True)
        txt(draw, (bx + 82, by + 65), sub, 15, "#617184")
    rounded(draw, (x, y + 455, x + w, y + 545), 18, "#FFF7E7", "#F59E0B")
    txt(draw, (x + 25, y + 483), "今日待办", 21, "#9D6205", True)
    txt(draw, (x + 25, y + 520), "合同待签 2　雇主险待增 3　离职交接 1", 20, "#5B4A2D")
    txt(draw, (1195, 340), "常用顺序", 26, "#182739", True)
    for i, line in enumerate(["先看待办", "再处理入职", "补合同/雇主险", "最后登记预支"]):
        draw.ellipse((1210, 400 + i * 78, 1250, 440 + i * 78), fill="#0EA5A8")
        txt(draw, (1230, 420 + i * 78), str(i + 1), 18, "#FFFFFF", True, "mm")
        txt(draw, (1270, 420 + i * 78), line, 23, "#395066", False, "lm")
    return save(image, "04_小程序驻厂工作台示意.png")


def make_roster():
    image, draw = canvas("小程序：驻厂人员管理", "按客户、生命周期、招聘渠道和雇主险状态筛选员工")
    x, y, w, h = mobile_frame(draw, "驻厂人员管理", "客户：常州精工制造")
    tabs = [("待入职", "12"), ("未入职", "4"), ("在职", "86"), ("离职中", "2")]
    for i, (label, value) in enumerate(tabs):
        bx = x + i * 146
        fill = "#0EA5A8" if i == 2 else "#EEF4F7"
        color = "#FFFFFF" if i == 2 else "#395066"
        rounded(draw, (bx, y, bx + 132, y + 70), 15, fill)
        txt(draw, (bx + 66, y + 23), label, 16, color, False, "mm")
        txt(draw, (bx + 66, y + 49), value, 21, color, True, "mm")
    rounded(draw, (x, y + 90, x + w, y + 145), 14, "#F3F7F9")
    txt(draw, (x + 22, y + 118), "⌕  姓名或手机号", 19, "#8393A3", False, "lm")
    employees = [
        ("张强", "常州精工制造", "普工", "已增保", "在职"),
        ("李敏", "常州精工制造", "质检员", "待增保", "待入职"),
        ("王磊", "武进电子项目", "装配工", "已增保", "离职中"),
    ]
    for i, row in enumerate(employees):
        by = y + 165 + i * 130
        rounded(draw, (x, by, x + w, by + 112), 18, "#FFFFFF", "#DCE8ED")
        draw.ellipse((x + 18, by + 18, x + 88, by + 88), fill="#DDF5F4")
        txt(draw, (x + 53, by + 53), row[0][0], 25, "#087D80", True, "mm")
        txt(draw, (x + 110, by + 26), row[0], 23, "#182739", True)
        txt(draw, (x + 110, by + 60), f"{row[1]} · {row[2]}", 18, "#617184")
        badge_color = "#16A36A" if row[3] == "已增保" else "#F59E0B"
        rounded(draw, (x + w - 205, by + 18, x + w - 105, by + 54), 12, badge_color)
        txt(draw, (x + w - 155, by + 36), row[3], 16, "#FFFFFF", True, "mm")
        rounded(draw, (x + w - 95, by + 18, x + w - 15, by + 54), 12, "#2563EB")
        txt(draw, (x + w - 55, by + 36), row[4], 16, "#FFFFFF", True, "mm")
        txt(draw, (x + w - 30, by + 82), "查看 ›", 18, "#0A8F93", True, "rm")
    txt(draw, (95, 300), "筛选维度", 28, "#182739", True)
    for i, line in enumerate(["客户/项目", "人员状态", "招聘渠道", "雇主险状态"]):
        rounded(draw, (95, 360 + i * 82, 345, 420 + i * 82), 16, "#FFFFFF", "#BFD8E2")
        txt(draw, (220, 390 + i * 82), line, 23, "#395066", True, "mm")
    return save(image, "05_小程序员工分类管理示意.png")


def make_add_employee():
    image, draw = canvas("新增员工：现场录入路径", "可手工填写，也可扫描身份证 OCR 自动带入姓名、身份证号和住址")
    steps = [
        ("1", "打开入口", "工作台 → 录入新员工"),
        ("2", "身份信息", "扫码身份证或手工填写"),
        ("3", "任职信息", "客户、项目、岗位（普工优先）"),
        ("4", "招聘信息", "招聘渠道关联供应商"),
        ("5", "保存档案", "默认待入职，现场再确认"),
    ]
    for i, (num, title, sub) in enumerate(steps):
        x = 65 + i * 305
        rounded(draw, (x, 245, x + 255, 465), 26, "#FFFFFF", "#BFD8E2", 2)
        draw.ellipse((x + 82, 275, x + 172, 365), fill="#0EA5A8")
        txt(draw, (x + 127, 320), num, 32, "#FFFFFF", True, "mm")
        txt(draw, (x + 127, 402), title, 27, "#182739", True, "mm")
        txt(draw, (x + 127, 440), sub, 17, "#617184", False, "mm")
        if i < 4:
            arrow(draw, (x + 262, 355), (x + 293, 355), width=5)
    rounded(draw, (190, 610, 1410, 785), 28, "#FFF7E7", "#F59E0B", 2)
    txt(draw, (240, 655), "保存前必查", 28, "#A66300", True)
    txt(draw, (240, 710), "姓名、身份证、手机号、客户、项目、岗位、用工模式、费用模式、招聘渠道", 27, "#453A28")
    txt(draw, (240, 755), "命中全公司黑名单时系统会拦截，不得绕过录入。", 24, "#B63A3A", True)
    return save(image, "06_新增员工操作流程.png")


def make_compliance():
    image, draw = canvas("新员工入职合规：只处理两件事", "确认入职后，驻厂专员及时补齐劳动合同和雇主险")
    blocks = [
        (120, "劳动合同", "登记合同信息\n选择“已签署”\n保留签署记录", "#2563EB", "合同待签"),
        (865, "雇主险", "选择“增保”\n填写生效信息\n离职时办理减保", "#16A36A", "雇主险待增"),
    ]
    for x, title, body, color, risk in blocks:
        rounded(draw, (x, 240, x + 615, 610), 32, "#FFFFFF", color, 3)
        rounded(draw, (x + 40, 280, x + 575, 355), 22, color)
        txt(draw, (x + 307, 318), title, 34, "#FFFFFF", True, "mm")
        for i, line in enumerate(body.split("\n")):
            draw.ellipse((x + 70, 410 + i * 55, x + 100, 440 + i * 55), fill=color)
            txt(draw, (x + 85, 425 + i * 55), "✓", 17, "#FFFFFF", True, "mm")
            txt(draw, (x + 125, 425 + i * 55), line, 25, "#395066", False, "lm")
        rounded(draw, (x + 175, 545, x + 440, 590), 14, "#FFF1F1" if color == "#2563EB" else "#ECF9F2")
        txt(draw, (x + 307, 568), f"未完成显示：{risk}", 18, "#B63A3A" if color == "#2563EB" else "#15724C", True, "mm")
    arrow(draw, (745, 425), (855, 425), color="#0EA5A8", width=7)
    rounded(draw, (300, 710, 1300, 800), 22, "#EAF7FA", "#0EA5A8")
    txt(draw, (800, 755), "两项均完成后，员工入职风险状态即为正常", 28, "#087D80", True, "mm")
    return save(image, "07_合同与雇主险合规图.png")


def make_advance():
    image, draw = canvas("驻厂预支登记", "登记预支时间、金额、用途和员工所属客户项目，形成可追溯台账")
    x, y, w, h = mobile_frame(draw, "驻厂预支记录", "现场登记后直接计入未结余额")
    fields = [
        ("在职员工 *", "张强"),
        ("客户/项目", "常州精工制造 / 一厂项目"),
        ("预支时间 *", "2026-08-10  14:30"),
        ("预支金额 *", "¥ 500.00"),
        ("预支用途 *", "生活费"),
    ]
    for i, (label, value) in enumerate(fields):
        by = y + i * 85
        txt(draw, (x, by + 12), label, 18, "#617184")
        rounded(draw, (x + 185, by - 5, x + w, by + 52), 13, "#F5F8FA", "#D9E5EA")
        txt(draw, (x + 205, by + 24), value, 19, "#182739", True, "lm")
    rounded(draw, (x, y + 440, x + w, y + 495), 15, "#0EA5A8")
    txt(draw, (x + w // 2, y + 468), "保存预支记录", 22, "#FFFFFF", True, "mm")
    rounded(draw, (120, 300, 390, 505), 24, "#FFFFFF", "#F59E0B", 2)
    txt(draw, (255, 345), "金额规则", 28, "#A66300", True, "mm")
    txt(draw, (255, 405), "单笔 ≤ 2,000 元", 24, "#453A28", True, "mm")
    txt(draw, (255, 455), "未结合计 ≤ 3,000 元", 24, "#453A28", True, "mm")
    txt(draw, (255, 495), "仅限在职员工", 21, "#B63A3A", True, "mm")
    rounded(draw, (1200, 300, 1490, 505), 24, "#FFFFFF", "#0EA5A8", 2)
    txt(draw, (1345, 345), "保存结果", 28, "#087D80", True, "mm")
    for i, line in enumerate(["生成预支台账", "增加未结余额", "薪资核算时扣回"]):
        txt(draw, (1235, 405 + i * 43), f"✓ {line}", 22, "#395066", True)
    return save(image, "08_驻厂预支登记示意.png")


def make_offboarding():
    image, draw = canvas("离职办理闭环", "发起离职不等于流程完成，必须继续完成交接、减保和工资结算")
    nodes = [
        ("发起离职", "填写离职日期、原因", "#DC4C4C"),
        ("现场交接", "工牌/工具/宿舍/考勤", "#F59E0B"),
        ("雇主险减保", "登记减保结果", "#0EA5A8"),
        ("工资结算", "核对预支与应付工资", "#2563EB"),
        ("正式离职", "权限失效并回流人才库", "#16A36A"),
    ]
    for i, (title, sub, color) in enumerate(nodes):
        x = 70 + i * 305
        rounded(draw, (x, 280, x + 250, 500), 28, "#FFFFFF", color, 3)
        txt(draw, (x + 125, 345), str(i + 1), 39, color, True, "mm")
        txt(draw, (x + 125, 410), title, 28, "#182739", True, "mm")
        txt(draw, (x + 125, 458), sub, 18, "#617184", False, "mm")
        if i < 4:
            arrow(draw, (x + 256, 390), (x + 295, 390), color="#3AAFB5", width=6)
    rounded(draw, (250, 650, 1350, 785), 24, "#FFF1F1", "#DC4C4C")
    txt(draw, (300, 695), "注意", 27, "#B63A3A", True)
    txt(draw, (300, 744), "离职员工不得继续使用在职业务权限；人员档案保留并自动进入人才库。", 27, "#5A3333")
    return save(image, "09_离职办理流程图.png")


def make_daily_checklist():
    image, draw = canvas("驻厂专员每日工作清单", "建议按“先风险、后人员、再台账”的顺序处理")
    rows = [
        ("上班后", "查看驻厂待办", "合同、雇主险、待入职、离职交接"),
        ("新员工到场", "录入并确认入职", "身份、客户、项目、岗位、渠道"),
        ("工作期间", "维护在职状态", "调岗、资料更正、雇主险变动"),
        ("发生预支", "即时登记台账", "时间、金额、用途、所属项目"),
        ("员工离场", "完成离职闭环", "交接、减保、工资结算、人才库"),
        ("下班前", "复核今日记录", "无遗漏、无跨项目、无错误状态"),
    ]
    for i, (time, action, check) in enumerate(rows):
        y = 185 + i * 105
        rounded(draw, (105, y, 1495, y + 82), 20, "#FFFFFF", "#D7E6EC")
        rounded(draw, (125, y + 14, 285, y + 68), 15, "#071B33")
        txt(draw, (205, y + 41), time, 21, "#FFFFFF", True, "mm")
        txt(draw, (330, y + 28), action, 24, "#182739", True)
        txt(draw, (650, y + 28), check, 22, "#617184")
        draw.ellipse((1425, y + 24, 1460, y + 59), outline="#0EA5A8", width=4)
    return save(image, "10_驻厂每日工作清单.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=95, start=110, bottom=95, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_keep(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def set_run(run, size=10.5, bold=False, color=TEXT):
    # LibreOffice 的无界面转换环境对部分 macOS 字体族识别不稳定；
    # Arial Unicode MS 可在 DOCX -> PDF 时完整保留中文，避免正文变成方框。
    name = "Arial Unicode MS"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", size=10.5, bold=False, color=TEXT, align=None, before=0, after=5, line=1.3):
    p = doc.add_paragraph()
    if text:
        set_run(p.add_run(text), size, bold, color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_run(p.add_run(text), 18 if level == 1 else 13.5, True, NAVY if level == 1 else NAVY_2)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(7)
    set_keep(p)
    return p


def add_steps(doc, steps):
    for i, step in enumerate(steps, 1):
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        t.columns[0].width = Cm(1.2)
        t.columns[1].width = Cm(15.2)
        left, right = t.rows[0].cells
        set_cell_shading(left, TEAL)
        set_cell_shading(right, PALE)
        left.vertical_alignment = right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        left.text = str(i)
        right.text = step
        for cell in (left, right):
            set_cell_margins(cell, 100, 130, 100, 130)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run(run, 10.5, cell is left, WHITE if cell is left else TEXT)
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph(doc, "", after=2)


def add_note(doc, title, text, tone="info"):
    colors = {
        "info": (LIGHT_BLUE, TEAL),
        "warn": ("FFF7E7", AMBER),
        "danger": ("FFF0F0", RED),
        "success": ("EAF8F1", GREEN),
    }
    fill, accent = colors[tone]
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    set_run(p.add_run(f"{title}　"), 10.5, True, accent)
    set_run(p.add_run(text), 10.5, False, TEXT)
    paragraph(doc, "", after=2)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item), 10.5, False, TEXT)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_repeat_header(t.rows[0])
    for i, h in enumerate(headers):
        t.cell(0, i).text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    for r_idx, row in enumerate(t.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, PALE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run(run, 9.2, r_idx == 0, WHITE if r_idx == 0 else TEXT)
    paragraph(doc, "", after=3)
    return t


def add_image(doc, path, caption, width=16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    c = paragraph(doc, f"图：{caption}", 9, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, after=8)
    return c


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run(run, 9, False, GRAY)


def page_break(doc):
    doc.add_page_break()


def build_doc(images):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    for style_name, size, color in (("Heading 1", 18, NAVY), ("Heading 2", 13.5, NAVY_2), ("Heading 3", 11.5, TEAL)):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = sec.header.paragraphs[0]
    set_run(header.add_run("优益数字化管理系统｜驻厂专员使用手册"), 8.5, False, GRAY)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(sec.footer.paragraphs[0])

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(images["cover"]), width=Cm(17.0))
    paragraph(doc, "内部培训资料｜请勿包含员工完整身份证号、银行卡号等敏感信息", 9, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, before=6)
    page_break(doc)

    add_heading(doc, "版本信息", 1)
    add_table(doc, ["项目", "内容"], [
        ["系统名称", "优益数字化管理系统"],
        ["手册名称", "驻厂专员使用手册"],
        ["版本", "V1.0"],
        ["发布日期", str(date.today())],
        ["适用终端", "微信小程序、手机 Web、电脑 Web"],
        ["适用角色", "已被企业管理员派遣到客户/项目的驻厂专员"],
    ], [4.0, 12.2])
    add_note(doc, "手册说明", "本手册配图为按当前系统功能制作的页面示意图。界面文字或位置随版本更新可能略有变化，以实际系统为准。", "info")

    add_heading(doc, "目录", 1)
    toc = [
        "1. 手册适用对象与工作目标", "2. 驻厂专员权限范围", "3. 登录与首次使用", "4. 驻厂工作台",
        "5. 查看负责客户与项目", "6. 新增员工", "7. 批量录入员工", "8. 扫描身份证 OCR",
        "9. 确认员工入职", "10. 在职员工维护与调岗", "11. 劳动合同与雇主险", "12. 驻厂预支登记",
        "13. 离职管理", "14. 人才库流转", "15. 驻厂待办", "16. 数据权限与保密",
        "17. 常见问题", "18. 每日工作清单", "附录：状态与字段说明",
    ]
    for item in toc:
        paragraph(doc, item, 10.5, False, TEXT, after=2)
    page_break(doc)

    add_heading(doc, "1. 手册适用对象与工作目标", 1)
    paragraph(doc, "本手册用于指导驻厂专员通过优益数字化管理系统管理本人负责客户/项目的员工。重点是让人员状态、合同、雇主险、预支和离职记录及时、准确、可追溯。")
    add_table(doc, ["工作目标", "驻厂要求"], [
        ["人员准确", "当天到场、未到场、在职、离职状态当天维护"],
        ["合规可控", "新员工入职后及时确认劳动合同和雇主险"],
        ["现场留痕", "预支、调岗、离职等事项在系统登记，不只在微信口头沟通"],
        ["数据隔离", "只处理本人授权客户/项目，不跨项目查看或导出员工资料"],
    ], [4.2, 12.0])
    add_image(doc, images["flow"], "驻厂员工全生命周期")

    add_heading(doc, "2. 驻厂专员权限范围", 1)
    paragraph(doc, "驻厂专员的数据范围由企业管理员配置。管理员在客户项目卡片中点击“派遣驻厂”，选择驻厂专员后，系统才会授予对应项目数据。")
    add_image(doc, images["permission"], "企业管理员派遣与驻厂数据权限关系")
    add_heading(doc, "2.1 可以操作", 2)
    add_bullets(doc, [
        "查看本人负责客户/项目及其员工。",
        "新增员工、批量录入员工、编辑本人项目员工。",
        "确认入职、调岗、办理离职。",
        "登记劳动合同、办理雇主险增保或减保。",
        "登记现场预支，查看本人范围内的预支台账。",
        "查看驻厂待办、项目人员状态和合规事项。",
    ])
    add_heading(doc, "2.2 不可以操作", 2)
    add_bullets(doc, [
        "查看未授权客户/项目的员工。",
        "修改其他驻厂专员负责项目的数据。",
        "绕过黑名单拦截录入员工。",
        "擅自扩大账号权限或为自己增加项目。",
        "在群聊、朋友圈或私人设备传播员工身份证、银行卡、工资等信息。",
    ])
    add_note(doc, "看不到客户怎么办", "先确认企业管理员是否已在项目卡点击“派遣驻厂”并选择你的账号。重新登录后仍看不到，再联系企业管理员检查账号状态和项目授权。", "warn")

    add_heading(doc, "3. 登录与首次使用", 1)
    add_heading(doc, "3.1 微信小程序登录", 2)
    add_steps(doc, [
        "打开微信，进入“优益数字化管理系统”小程序。",
        "输入企业分配的账号和密码，点击登录。",
        "首次登录后进入“我的”，确认姓名、角色显示为驻厂专员。",
        "进入“驻厂”页，检查是否能看到本人负责客户和项目。",
    ])
    add_heading(doc, "3.2 网页端登录", 2)
    add_steps(doc, [
        "在手机或电脑浏览器打开公司提供的系统网址。",
        "输入驻厂账号和密码，点击登录。",
        "进入办公中心或员工花名册，核对客户/项目范围。",
    ])
    add_note(doc, "账号安全", "不要共用账号，不要把验证码、密码或登录 Token 转发给他人；离开电脑时退出登录。", "danger")

    add_heading(doc, "4. 驻厂工作台", 1)
    paragraph(doc, "工作台是驻厂专员每天的第一入口，汇总在职人数、在营项目、未结预支、合规待办和员工生命周期入口。")
    add_image(doc, images["workbench"], "小程序驻厂工作台页面示意")
    add_table(doc, ["区域", "用途", "建议动作"], [
        ["核心指标", "在职人员、在营项目、未结预支", "发现数量异常时进入对应列表核对"],
        ["驻厂人员管理", "录入、待入职、在职、离职交接", "按员工当前阶段处理"],
        ["驻厂处理队列", "合同、雇主险、转岗、离职事项", "优先处理红色/逾期待办"],
        ["现场快捷处理", "新增员工、我的员工、招聘渠道、预支", "高频操作直接进入"],
    ], [3.4, 6.0, 6.8])

    add_heading(doc, "5. 查看负责客户与项目", 1)
    add_steps(doc, [
        "在小程序底部点击“驻厂”。",
        "顶部客户分类中选择客户单位；列表仅显示授权范围。",
        "查看客户下待入职、未入职、在职、离职中和雇主险待增人数。",
        "再按人员状态、招聘渠道或雇主险状态筛选。",
        "点击员工卡片进入详情，查看和处理个人事项。",
    ])
    add_image(doc, images["roster"], "员工按客户、状态和雇主险分类管理")
    add_note(doc, "项目关联规则", "员工必须关联正确客户和项目。驻厂专员只能在本人授权范围内选择项目；若项目为空，先联系企业管理员检查客户项目和派遣配置。", "warn")

    add_heading(doc, "6. 新增员工", 1)
    paragraph(doc, "适用于新员工到场前预录入或现场报到。员工默认保存为“待入职”，确认实际到岗后再执行“确认入职”。")
    add_image(doc, images["add"], "新增员工录入流程")
    add_heading(doc, "6.1 操作步骤", 2)
    add_steps(doc, [
        "工作台点击“录入新员工”，或驻厂人员管理右上角点击“录入员工”。",
        "选择扫描身份证或手工填写姓名、身份证号、手机号。",
        "选择客户单位和所属项目。只可选择本人负责范围。",
        "选择岗位。岗位列表第一项为“普工”，也可选择装配工、操作工、质检员、仓管员、文员等。",
        "填写用工模式；费用模式不设默认值，按实际合作约定填写。",
        "填写招聘渠道。招聘渠道可关联渠道商/供应商，不再分别填写招聘类型和招聘人。",
        "录入状态选择“待入职”“直接入职”或“未入职”。一般现场预录选择“待入职”。",
        "核对无误后保存。系统会校验重复身份证、手机号和全公司黑名单。",
    ])
    add_heading(doc, "6.2 必填及关键字段", 2)
    add_table(doc, ["字段", "填写规则", "示例"], [
        ["姓名", "与身份证一致", "张强"],
        ["身份证号码", "18位；不得截图外传", "系统录入后脱敏显示"],
        ["手机号", "员工本人常用号码", "138****8000"],
        ["客户单位/项目", "必须选择实际到岗单位和项目", "常州精工制造/一厂项目"],
        ["岗位", "按实际岗位；普工排在第一位", "普工"],
        ["用工模式", "按项目合同选择", "派遣/外包/灵活用工"],
        ["费用模式", "空白自行填写，不自动默认", "管理费600元/月"],
        ["招聘渠道", "填写并关联渠道商/供应商", "XX人力供应商"],
        ["录入状态", "待入职/直接入职/未入职", "待入职"],
    ], [3.3, 8.3, 4.6])

    add_heading(doc, "7. 批量录入员工", 1)
    paragraph(doc, "适合集中到岗或从 Excel 名单导入。单次最多 200 人。")
    add_steps(doc, [
        "网页端进入员工花名册，点击“批量录入”。",
        "下载 CSV 或 XLSX 模板，不要修改表头名称和列顺序。",
        "填写员工信息；用工模式、费用模式必填。费用模式按实际填写。",
        "录入状态可填待入职、直接入职、未入职；留空默认待入职。",
        "上传文件，或将 Excel 内容复制到粘贴区域。",
        "执行导入后查看成功和失败数量；失败行按提示修正后单独重传。",
    ])
    add_note(doc, "导入前检查", "客户单位、项目、岗位和招聘渠道应使用系统现有或允许的名称；同一身份证、手机号重复或命中黑名单会导入失败。", "warn")

    add_heading(doc, "8. 扫描身份证 OCR", 1)
    add_steps(doc, [
        "进入新增员工页面，点击“扫描身份证”。",
        "允许小程序使用相机，拍摄身份证人像面。",
        "保持证件完整、光线均匀、文字清晰，不要反光或遮挡。",
        "识别完成后核对姓名、身份证号码、性别、住址。",
        "OCR 只负责减少手工输入，驻厂专员仍需核对原件。",
    ])
    add_note(doc, "识别失败", "重新拍摄并保证证件四角完整；仍失败时改为手工录入。禁止把身份证照片发到工作群请求他人代录。", "danger")

    add_heading(doc, "9. 确认员工入职", 1)
    paragraph(doc, "“待入职”表示已建档但尚未确认到岗。员工实际到厂并通过现场核验后，才能确认入职。")
    add_steps(doc, [
        "在驻厂人员管理选择“待入职”。",
        "找到员工，点击“确认入职”。",
        "复核客户、项目、岗位、入职日期和招聘渠道。",
        "确认员工已实际到岗后提交，状态变为“在职”。",
        "继续办理劳动合同和雇主险增保，完成入职合规。",
    ])
    add_note(doc, "未到岗处理", "员工最终未到岗时，将状态改为“未入职”。其信息会进入人才库继续跟进，不应误设为在职。", "info")

    add_heading(doc, "10. 在职员工维护与调岗", 1)
    add_heading(doc, "10.1 编辑员工", 2)
    add_steps(doc, [
        "进入驻厂人员管理，选择客户后找到员工。",
        "点击员工卡片，再点击“编辑资料”。",
        "修改手机号、岗位、费用模式、招聘渠道等允许字段。",
        "保存后返回列表，确认展示已更新。",
    ])
    add_heading(doc, "10.2 员工调岗", 2)
    add_steps(doc, [
        "从员工详情发起调岗，选择目标项目、岗位和生效日期。",
        "同项目或授权范围内调岗按系统提示确认。",
        "跨项目调岗由目标项目驻厂专员接收；跨客户调岗需企业 HR 复核。",
        "接收完成后，新项目驻厂专员可在本人客户/项目下看到员工。",
    ])
    add_note(doc, "不要直接改历史", "调岗必须使用“调岗”功能，系统才会保留任职记录。不要仅修改岗位名称代替调岗。", "warn")

    add_heading(doc, "11. 劳动合同与雇主险", 1)
    paragraph(doc, "用工风险中心保持简化：新员工入职主要检查劳动合同是否签订、雇主险是否增保。")
    add_image(doc, images["compliance"], "劳动合同与雇主险两项入职合规")
    add_heading(doc, "11.1 登记劳动合同", 2)
    add_steps(doc, [
        "打开在职员工详情，点击“登记合同”。",
        "填写合同类型、开始日期、结束日期和签署状态。",
        "只有实际完成签署后才能选择“已签署”。",
        "保存后回到员工详情，确认合同状态已更新。",
    ])
    add_heading(doc, "11.2 雇主险增保", 2)
    add_steps(doc, [
        "打开在职员工详情，点击“雇主险”。",
        "选择“增保”，填写生效日期和必要备注。",
        "保存后确认员工卡片显示“已增保/保障中”。",
    ])
    add_heading(doc, "11.3 雇主险减保", 2)
    add_steps(doc, [
        "员工发起离职后，从离职交接或员工详情进入雇主险。",
        "选择“减保”，填写减保日期和备注。",
        "保存后继续完成工资结算和其他离职交接。",
    ])
    add_note(doc, "合规范围", "当前系统不再设置独立保险提示、社保公积金管理，只保留雇主险增保/减保。", "info")

    add_heading(doc, "12. 驻厂预支登记", 1)
    paragraph(doc, "驻厂预支用于记录员工现场实际预支情况。登记后直接形成台账和未结余额，薪资核算时按记录扣回。")
    add_image(doc, images["advance"], "小程序驻厂预支登记页面示意")
    add_steps(doc, [
        "在工作台点击“驻厂预支”，点击“登记预支”。",
        "选择在职员工。系统自动带出客户单位和项目。",
        "填写实际预支日期、时间、金额和用途。",
        "核对后保存；系统记录登记人和操作时间。",
        "在预支台账中确认金额、用途、客户项目和未结余额。",
    ])
    add_table(doc, ["校验项", "规则"], [
        ["员工状态", "仅在职员工可登记"],
        ["单笔金额", "大于0且不超过2,000元"],
        ["未结合计", "员工未结余额加本次金额不超过3,000元"],
        ["时间范围", "不可晚于当前时间，仅允许补录最近一年"],
        ["所属项目", "驻厂专员账号必须选择本人授权项目"],
    ], [5.0, 11.2])

    add_heading(doc, "13. 离职管理", 1)
    add_image(doc, images["offboarding"], "员工离职办理闭环")
    add_steps(doc, [
        "在在职员工卡片点击“办理离职”。",
        "填写离职日期、离职原因和备注，发起离职。",
        "按现场实际完成工牌、工具、宿舍、考勤等交接。",
        "办理雇主险减保。",
        "核对工资、预支扣回和结算状态。",
        "所有事项完成后提交，员工转为“离职”。",
    ])
    add_note(doc, "状态说明", "发起离职后员工处于“离职中”，仍需继续处理。只有交接、减保和工资结算闭环后才成为正式离职。", "warn")

    add_heading(doc, "14. 人才库流转", 1)
    paragraph(doc, "人才库与新增员工、花名册关联，用于沉淀未入职和已离职人员。")
    add_table(doc, ["触发状态", "系统动作", "驻厂处理"], [
        ["未入职", "保留员工档案并自动进入人才库", "补充未到岗原因和后续跟进"],
        ["离职", "正式离职后自动回流人才库", "保留原客户、岗位、渠道和离职信息"],
        ["重新录用", "从人才库/未入职人员重新发起录用", "重新核对客户、项目、岗位和入职合规"],
    ], [3.0, 6.0, 7.2])
    add_note(doc, "禁止重复建档", "已有员工或人才记录时优先使用重新录用或关联功能，不要新建重复人员。", "danger")

    add_heading(doc, "15. 驻厂待办", 1)
    paragraph(doc, "驻厂待办按当前账号授权项目统计，常见事项包括待入职、合同待签、雇主险待增、转岗接收、离职减保和离职交接。")
    add_steps(doc, [
        "每天登录后先打开驻厂处理队列。",
        "点击待办进入对应员工或业务页面。",
        "完成业务操作后返回工作台，系统会刷新数量。",
        "若待办未消失，检查是否还有必要字段或流程节点未完成。",
    ])
    add_table(doc, ["待办", "处理入口", "完成标准"], [
        ["待确认入职", "待入职人员 → 确认入职", "员工状态为在职"],
        ["合同待签", "员工详情 → 登记合同", "签署状态为已签署"],
        ["雇主险待增", "员工详情 → 雇主险 → 增保", "状态为保障中"],
        ["转岗待接收", "驻厂待办 → 转岗记录", "接收或按实际拒绝"],
        ["离职待减保", "离职交接 → 雇主险减保", "减保记录已保存"],
        ["离职交接", "员工详情 → 继续离职交接", "交接和工资结算全部完成"],
    ], [3.6, 6.2, 6.4])

    add_heading(doc, "16. 数据权限与保密", 1)
    add_bullets(doc, [
        "员工身份证、手机号、银行卡、工资属于敏感个人信息，只能用于授权业务。",
        "系统默认脱敏显示时，不得通过其他方式索取并传播完整信息。",
        "禁止将名单导出到私人电脑、私人网盘或无权限工作群。",
        "禁止代替他人登录；账号变更、离职或调离项目时应立即通知管理员调整权限。",
        "发现查看到非本人项目数据时停止操作，立即报告企业管理员。",
        "黑名单信息全公司共享，但不得用于与业务无关的传播或歧视性用途。",
    ])
    add_note(doc, "操作留痕", "员工新增、编辑、合同、雇主险、预支、调岗和离职等关键操作会记录操作人和时间。请使用本人账号据实操作。", "info")

    add_heading(doc, "17. 常见问题", 1)
    add_table(doc, ["问题", "可能原因", "处理方法"], [
        ["登录后没有提示或页面空白", "网络、账号停用或部分数据加载失败", "检查网络并重新登录；仍失败联系管理员提供账号和发生时间"],
        ["看不到客户/项目", "管理员未派遣或授权已失效", "请管理员在项目卡执行“派遣驻厂”"],
        ["新增员工按钮不显示", "账号缺少员工新增权限", "请企业管理员配置驻厂专员的员工新增/批量录入权限"],
        ["选择不到项目", "账号未授权该项目", "先完成项目派遣，不要选择其他客户项目替代"],
        ["身份证 OCR 失败", "反光、裁切、模糊或 OCR 服务异常", "重新拍摄；仍失败时手工录入并核验原件"],
        ["员工保存失败", "必填项缺失、重复人员或命中黑名单", "按系统提示补全或停止录入并报告"],
        ["费用模式无法选择", "费用模式设计为空白自行填写", "按客户合同实际约定填写"],
        ["预支无法保存", "非在职、超单笔/累计额度、项目未关联", "核对状态、金额和所属项目"],
        ["离职后仍显示离职中", "交接、减保或工资结算未完成", "进入继续离职交接，逐项完成"],
        ["操作后数量未更新", "页面缓存或接口未刷新", "返回工作台下拉刷新或重新进入页面"],
    ], [4.0, 5.4, 6.8])

    add_heading(doc, "18. 每日工作清单", 1)
    add_image(doc, images["daily"], "驻厂专员每日工作清单")
    add_table(doc, ["检查项", "完成确认"], [
        ["今日到场新员工均已录入，客户/项目/岗位正确", "□"],
        ["待入职、未入职、在职、离职状态与现场一致", "□"],
        ["新入职员工合同和雇主险已处理或有明确待办", "□"],
        ["今日预支均已登记时间、金额、用途和项目", "□"],
        ["离职人员已继续处理交接、减保和工资结算", "□"],
        ["没有在聊天群发送完整身份证、银行卡或工资信息", "□"],
    ], [14.5, 1.7])

    add_heading(doc, "附录 A：员工状态说明", 1)
    add_table(doc, ["状态值", "状态名称", "使用场景", "后续动作"], [
        ["1", "待入职", "已建档，尚未确认实际到岗", "到岗后确认入职；未到岗改未入职"],
        ["2", "在职", "已到岗并进入在职管理", "合同、雇主险、调岗、预支"],
        ["3", "离职", "离职闭环完成", "档案保留并回流人才库"],
        ["4", "黑名单", "存在公司级风险记录", "禁止绕过拦截录用"],
        ["5", "未入职", "未到岗或未完成录用", "自动进入人才库继续跟进"],
    ], [2.0, 3.0, 6.0, 5.2])

    add_heading(doc, "附录 B：问题反馈模板", 1)
    paragraph(doc, "向管理员反馈系统问题时，请复制以下内容填写。不要附员工完整身份证号、银行卡号或密码。")
    add_table(doc, ["项目", "填写内容"], [
        ["发生时间", "年/月/日 时:分"],
        ["使用终端", "微信小程序 / 手机网页 / 电脑网页"],
        ["账号角色", "驻厂专员"],
        ["客户/项目", "填写名称，不发送敏感个人信息"],
        ["操作路径", "例如：驻厂 → 在职员工 → 张* → 雇主险"],
        ["实际结果", "系统提示或页面现象"],
        ["期望结果", "希望完成的业务动作"],
        ["截图", "可截页面；先遮挡身份证、手机号、银行卡、工资"],
    ], [4.0, 12.2])
    add_note(doc, "培训完成标准", "驻厂专员应能独立完成：新增员工 → 确认入职 → 登记合同/雇主险 → 登记预支 → 办理离职闭环。", "success")

    # 新章节的首页取消与上一页相同，保持页眉页脚统一。
    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        "cover": make_cover(),
        "permission": make_permission(),
        "flow": make_flow(),
        "workbench": make_workbench(),
        "roster": make_roster(),
        "add": make_add_employee(),
        "compliance": make_compliance(),
        "advance": make_advance(),
        "offboarding": make_offboarding(),
        "daily": make_daily_checklist(),
    }
    build_doc(images)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
